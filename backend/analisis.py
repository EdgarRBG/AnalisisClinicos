import sqlite3
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

def conectar():
    return sqlite3.connect("sistema_laboratorio.db")

class Analisis:
    def registrar_solicitud(self, paciente_id, cita_id=None, id_muestra=None, medico_solicitante="", tipo_estudio="", observaciones=""):
        conn = conectar()
        cursor = conn.cursor()
        try:
            if not paciente_id:
                raise ValueError("Paciente es obligatorio para crear solicitud")

            fecha_recepcion = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            if not id_muestra:
                id_muestra = datetime.now().strftime("%Y%m%d%H%M%S")
                print(f"ID Muestra generado automáticamente: {id_muestra}")

            cursor.execute("""
                INSERT INTO solicitudes_analisis 
                (paciente_id, cita_id, id_muestra, medico_solicitante, tipo_estudio, observaciones_generales, fecha_recepcion, estado)
                VALUES (?, ?, ?, ?, ?, ?, ?, 'pendiente')
            """, (paciente_id, cita_id, id_muestra, medico_solicitante, tipo_estudio, observaciones, fecha_recepcion))
            
            conn.commit()
            print(f"Solicitud creada con ID: {cursor.lastrowid}")
            return cursor.lastrowid
        except Exception as e:
            print("Error al registrar solicitud:", e)
            return None
        finally:
            conn.close()

    def agregar_resultado(self, solicitud_id, parametro, resultado=None, unidades="", valor_referencia="", fuera_de_rango="", observacion=""):
        conn = conectar()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                INSERT INTO resultados_parametros 
                (solicitud_id, parametro, resultado, unidades, valor_referencia, fuera_de_rango, observacion)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (solicitud_id, parametro, resultado, unidades, valor_referencia, fuera_de_rango, observacion))
            conn.commit()
            print(f"Resultado agregado: {parametro} = {resultado} {unidades}")
            return True
        except Exception as e:
            print("Error al agregar resultado:", e)
            return False
        finally:
            conn.close()

    def obtener_pendientes(self):
        conn = conectar()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                SELECT 
                    s.id,
                    COALESCE(p.nombre, 'Paciente no registrado') AS paciente_nombre,
                    s.id_muestra,
                    s.tipo_estudio,
                    s.fecha_recepcion
                FROM solicitudes_analisis s
                LEFT JOIN pacientes p ON s.paciente_id = p.id
                WHERE s.estado IN ('pendiente', 'en proceso')
                ORDER BY s.fecha_recepcion DESC
            """)
            rows = cursor.fetchall()
            print(f"Pendientes encontradas: {len(rows)} solicitudes")
            return [dict(zip(['id', 'paciente_nombre', 'id_muestra', 'tipo_estudio', 'fecha_recepcion'], row)) for row in rows]
        except Exception as e:
            print("Error al obtener pendientes:", str(e))
            return []
        finally:
            conn.close()

    def obtener_solicitud(self, solicitud_id):
        conn = conectar()
        cursor = conn.cursor()
        try:
            print(f"Obteniendo solicitud ID: {solicitud_id}")
            cursor.execute("""
                SELECT 
                    s.id,
                    s.paciente_id,
                    s.id_muestra,
                    s.tipo_estudio,
                    s.fecha_recepcion,
                    s.fecha_reporte,
                    s.medico_solicitante,
                    s.estado,
                    s.observaciones_generales,
                    p.nombre AS paciente_nombre,
                    p.edad,
                    p.sexo
                FROM solicitudes_analisis s
                LEFT JOIN pacientes p ON s.paciente_id = p.id
                WHERE s.id = ?
            """, (solicitud_id,))
            solicitud = cursor.fetchone()

            if not solicitud:
                print(f"No se encontró solicitud con ID {solicitud_id}")
                return None

            solicitud_dict = {
                'id': solicitud[0],
                'paciente_id': solicitud[1],
                'id_muestra': solicitud[2],
                'tipo_estudio': solicitud[3],
                'fecha_recepcion': solicitud[4],
                'fecha_reporte': solicitud[5],
                'medico_solicitante': solicitud[6],
                'estado': solicitud[7],
                'observaciones_generales': solicitud[8],
                'paciente_nombre': solicitud[9],
                'edad': solicitud[10],
                'sexo': solicitud[11]
            }

            print("Solicitud encontrada:", solicitud_dict)

            cursor.execute("""
                SELECT parametro, resultado, unidades, valor_referencia, observacion
                FROM resultados_parametros 
                WHERE solicitud_id = ? 
                ORDER BY parametro
            """, (solicitud_id,))
            resultados = cursor.fetchall()

            print(f"Encontrados {len(resultados)} resultados")
            for r in resultados:
                print(f" - {r[0]}: resultado={r[1]}, unidades={r[2]}")

            return {
                "solicitud": solicitud_dict,
                "resultados": [dict(zip(['parametro', 'resultado', 'unidades', 'valor_referencia', 'observacion'], r)) for r in resultados]
            }
        except Exception as e:
            print("Error al obtener solicitud:", str(e))
            return None
        finally:
            conn.close()

    def actualizar_estado(self, solicitud_id, nuevo_estado):
        conn = conectar()
        cursor = conn.cursor()
        try:
            fecha_reporte = datetime.now().strftime("%Y-%m-%d %H:%M:%S") if nuevo_estado == "finalizado" else None
            cursor.execute("""
                UPDATE solicitudes_analisis 
                SET estado = ?, fecha_reporte = ?
                WHERE id = ?
            """, (nuevo_estado, fecha_reporte, solicitud_id))
            conn.commit()
            return cursor.rowcount > 0
        except Exception as e:
            print("Error al actualizar estado:", e)
            return False
        finally:
            conn.close()

    def obtener_finalizadas(self):
        conn = conectar()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                SELECT 
                    s.id,
                    COALESCE(p.nombre, 'Paciente no registrado') AS paciente_nombre,
                    s.id_muestra,
                    s.tipo_estudio,
                    s.fecha_recepcion,
                    s.fecha_reporte
                FROM solicitudes_analisis s
                LEFT JOIN pacientes p ON s.paciente_id = p.id
                WHERE s.estado = 'finalizado'
                ORDER BY s.fecha_reporte DESC
            """)
            rows = cursor.fetchall()
            return [dict(zip(['id', 'paciente_nombre', 'id_muestra', 'tipo_estudio', 'fecha_recepcion', 'fecha_reporte'], row)) for row in rows]
        except Exception as e:
            print("Error al obtener finalizadas:", e)
            return []
        finally:
            conn.close()

    def generar_reporte(self, solicitud_id):
        conn = conectar()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                SELECT 
                    s.id,
                    s.paciente_id,
                    s.id_muestra,
                    s.tipo_estudio,
                    s.fecha_recepcion,
                    s.fecha_reporte,
                    s.medico_solicitante,
                    s.estado,
                    s.observaciones_generales,
                    p.nombre AS paciente_nombre,
                    p.edad,
                    p.sexo
                FROM solicitudes_analisis s
                LEFT JOIN pacientes p ON s.paciente_id = p.id
                WHERE s.id = ?
            """, (solicitud_id,))
            solicitud = cursor.fetchone()

            if not solicitud:
                print(f"No se encontró solicitud para reporte: {solicitud_id}")
                return None

            cursor.execute("""
                SELECT parametro, resultado, unidades, valor_referencia, observacion
                FROM resultados_parametros 
                WHERE solicitud_id = ? 
                ORDER BY parametro
            """, (solicitud_id,))
            resultados = cursor.fetchall()

            doc = Document("combined_template.docx")

            for paragraph in doc.paragraphs:
                if "«PACIENTE»" in paragraph.text:
                    paragraph.text = paragraph.text.replace("«PACIENTE»", solicitud['paciente_nombre'] or "—")
                if "«EDAD»" in paragraph.text:
                    paragraph.text = paragraph.text.replace("«EDAD»", str(solicitud['edad']) or "—")
                if "«SEXO»" in paragraph.text:
                    paragraph.text = paragraph.text.replace("«SEXO»", solicitud['sexo'] or "—")
                if "«FECHA_RECEPCION»" in paragraph.text:
                    paragraph.text = paragraph.text.replace("«FECHA_RECEPCION»", solicitud['fecha_recepcion'] or "—")
                if "«FECHA_REPORTE»" in paragraph.text:
                    paragraph.text = paragraph.text.replace("«FECHA_REPORTE»", solicitud['fecha_reporte'] or "—")
                if "«MEDICO»" in paragraph.text:
                    paragraph.text = paragraph.text.replace("«MEDICO»", solicitud['medico_solicitante'] or "—")
                if "«TIPO_ESTUDIO»" in paragraph.text:
                    paragraph.text = paragraph.text.replace("«TIPO_ESTUDIO»", solicitud['tipo_estudio'] or "—")

            for table in doc.tables:
                if "Resultado" in table.cell(0,0).text:
                    for res in resultados:
                        row = table.add_row().cells
                        row[0].text = res['parametro']
                        row[1].text = str(res['resultado']) if res['resultado'] is not None else "—"
                        row[2].text = res['unidades'] or "—"
                        row[3].text = res['valor_referencia'] or "—"
                        row[4].text = res['observacion'] or "—"

            output_filename = f"reporte_solicitud_{solicitud_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(output_filename)
            return output_filename

        except Exception as e:
            print("Error al generar reporte:", e)
            return None
        finally:
            conn.close()

    def importarDesdeCSV(self, solicitud_id, csv_content):
        import io
        import csv

        conn = conectar()
        cursor = conn.cursor()

        try:
            cursor.execute("SELECT id FROM solicitudes_analisis WHERE id = ? AND estado IN ('pendiente', 'en proceso')", (solicitud_id,))
            if not cursor.fetchone():
                return {"success": False, "error": "Solicitud no encontrada o ya finalizada"}

            csv_file = io.StringIO(csv_content)
            reader = csv.DictReader(csv_file)

            reader.fieldnames = [name.strip().upper() if name else '' for name in reader.fieldnames]
            print("Encabezados detectados en el CSV:", reader.fieldnames)

            mapeo = {
                'GLU': 'Glucosa',
                'GLUK': 'Glucosa',
                'CREAT': 'Creatinina',
                'UREA': 'Urea',
                'TRIG': 'Triglicéridos',
                'COLES': 'Colesterol Total',
                'TGP': 'Alanina Aminotransferasa (ALT)',
                'TGO': 'Aspartato Aminotransferasa (AST)',
                'A.U': 'Ácido Úrico',
                'A.U VTK': 'Ácido Úrico',
                'P.TOT': 'Proteínas Totales',
                'ALB': 'Albúmina',
                'FK': 'Fósforo',
                'K': 'Potasio',
                'Ca': 'Calcio',
                'Na': 'Sodio',
                'Cl': 'Cloro',
                'HDL BS': 'HDL Colesterol',
                'BD': 'Bilirrubina Directa',
                'BT': 'Bilirrubina Total',
                'He': 'Hemoglobina',
                'PCR': 'Proteína C Reactiva',
                'GGT': 'Gamma Glutamil Transferasa (GGT)',
                'LDH ELI': 'LDH',
                'AMI': 'Amilasa',
            }

            count = 0
            for row in reader:
                prueba_key = next((k for k in row if 'PRUEBA' in k.upper()), None)
                resultado_key = next((k for k in row if 'RESULTAD' in k.upper()), None)

                if not prueba_key or not resultado_key:
                    print("Fila sin Prueba o Resultados:", row)
                    continue

                prueba = row[prueba_key].strip().upper()
                resultado_str = row[resultado_key].strip()

                print(f"Prueba: '{prueba}', Resultado crudo: '{resultado_str}'")

                if not prueba or not resultado_str or resultado_str in ['0', '']:
                    print("Fila ignorada:", row)
                    continue

                resultado_str = resultado_str.replace(',', '.')  
                try:
                    resultado = float(resultado_str)
                except ValueError:
                    resultado = None
                    print("Resultado no numérico:", resultado_str)

                nombre_param = mapeo.get(prueba, prueba)

                unidades = 'mg/dL' if any(x in prueba for x in ['GLU', 'UREA', 'CREAT', 'COLES', 'TRIG']) else ''

                cursor.execute("""
                    INSERT INTO resultados_parametros 
                    (solicitud_id, parametro, resultado, unidades, valor_referencia, observacion)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (solicitud_id, nombre_param, resultado, unidades, '', 'Importado desde máquina'))

                count += 1
                print(f"Agregado: {nombre_param} = {resultado} {unidades}")

            conn.commit()
            print(f"Total parámetros agregados: {count}")
            return {"success": True, "count": count}

        except Exception as e:
            print("Error al importar CSV:", str(e))
            conn.rollback()
            return {"success": False, "error": str(e)}

        finally:
            conn.close()

    def eliminar_solicitud(self, solicitud_id):
        conn = conectar()
        cursor = conn.cursor()
        try:
            # Primero eliminamos los resultados asociados
            cursor.execute("DELETE FROM resultados_parametros WHERE solicitud_id = ?", (solicitud_id,))
            
            # Luego eliminamos la solicitud
            cursor.execute("DELETE FROM solicitudes_analisis WHERE id = ?", (solicitud_id,))
            
            filas_afectadas = cursor.rowcount
            conn.commit()
            
            print(f"Solicitud {solicitud_id} eliminada → {filas_afectadas} fila(s) afectada(s) en solicitudes_analisis")
            return filas_afectadas > 0
            
        except Exception as e:
            print("Error al eliminar solicitud:", str(e))
            conn.rollback()
            return False
        finally:
            conn.close()
